#!/usr/bin/env python3
"""Render an ATS-safe .docx and a plain-text .txt from a structured markdown CV.

ATS-safe by construction: single column, no tables, no text boxes, no images,
no headers/footers, standard fonts, conventional section headings. This is a
hold-the-line invariant of the skill — do not add tables or columns here.

Input markdown convention (produced by the tailorer):

    # Full Name
    City, Country | +44 7… | email@x.com | linkedin.com/in/handle | github.com/handle

    ## Professional Summary
    A short paragraph.

    ## Experience
    ### Job Title — Organisation, City
    September 2023 – July 2024
    - Bullet with verb + task + method + outcome.
    - Another bullet.

    ## Education
    ### MSc Something — University, City
    2024
    - Optional detail line.

    ## Skills
    - Category: item, item, item

    ## Certifications
    - Item

Rules:
  #    -> name (document title line, large bold)
  next non-empty, non-## line after the name -> contact line
  ##   -> section heading
  ###  -> entry sub-heading (role / degree title)
  -    -> bullet (level 1); "  -" (indented) -> bullet level 2
  other non-empty line -> normal paragraph (dates, summary text)

Usage:
  python render_docx.py --in cv.md --out-docx CV.docx --out-txt CV.txt
  python render_docx.py --in cv.md --outdir <folder>   # writes CV.docx + CV.txt
"""
import argparse
import os
import re
import sys

# UTF-8 stdout/stderr so printing a CV path or heading with accents/£/em-dashes never
# dies with a UnicodeEncodeError on a default (cp1252) Windows console.
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except (AttributeError, ValueError):
    pass

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Mm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("python-docx is required: pip install python-docx")

# Page size by market. UK-first default is A4; Canada/US use US Letter. python-docx
# otherwise defaults to Letter, so UK renders needed this to be correct.
PAGE_SIZES = {
    "a4": (Mm(210), Mm(297)),         # UK and most of the world
    "letter": (Inches(8.5), Inches(11)),  # Canada, US
}

BODY_FONT = "Calibri"      # widely parseable, standard
BODY_SIZE = 10.5
NAME_SIZE = 20
SECTION_SIZE = 12
ENTRY_SIZE = 11
DARK = RGBColor(0x1F, 0x1F, 0x1F)
ACCENT = RGBColor(0x30, 0x54, 0x96)


def parse(md):
    """Return a structured list of ('kind', text/level) tokens."""
    lines = md.replace("\r\n", "\n").split("\n")
    tokens = []
    name_seen = False
    contact_pending = False
    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            if contact_pending:
                contact_pending = False
            continue
        if stripped.startswith("# ") and not stripped.startswith("## "):
            tokens.append(("name", stripped[2:].strip()))
            name_seen = True
            contact_pending = True
            continue
        if contact_pending:
            tokens.append(("contact", stripped))
            contact_pending = False
            continue
        if stripped.startswith("### "):
            tokens.append(("entry", stripped[4:].strip()))
            continue
        if stripped.startswith("## "):
            tokens.append(("section", stripped[3:].strip()))
            continue
        m = re.match(r"^(\s*)-\s+(.*)$", raw)
        if m:
            indent = len(m.group(1))
            level = 2 if indent >= 2 else 1
            tokens.append(("bullet", (level, m.group(2).strip())))
            continue
        tokens.append(("para", stripped))
    if not name_seen:
        sys.exit("Input has no '# Name' line — cannot render a CV.")
    return tokens


def _strip_md_inline(text):
    # keep it simple: drop ** and * emphasis markers; ATS reads plain text best
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text


def build_docx(tokens, out_path, page="a4"):
    doc = Document()
    # base style
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = DARK
    width, height = PAGE_SIZES.get(page, PAGE_SIZES["a4"])
    # tight margins for two-page fit; explicit page size (A4 vs US Letter by market)
    for s in doc.sections:
        s.page_width, s.page_height = width, height
        s.top_margin = s.bottom_margin = Pt(36)      # 0.5"
        s.left_margin = s.right_margin = Pt(54)      # 0.75"

    def spacer(before=0, after=0):
        pass

    for kind, val in tokens:
        if kind == "name":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(_strip_md_inline(val))
            r.bold = True
            r.font.size = Pt(NAME_SIZE)
            r.font.color.rgb = DARK
            p.paragraph_format.space_after = Pt(2)
        elif kind == "contact":
            p = doc.add_paragraph()
            r = p.add_run(_strip_md_inline(val))
            r.font.size = Pt(BODY_SIZE)
            r.font.color.rgb = DARK
            p.paragraph_format.space_after = Pt(8)
        elif kind == "section":
            p = doc.add_paragraph()
            r = p.add_run(_strip_md_inline(val).upper())
            r.bold = True
            r.font.size = Pt(SECTION_SIZE)
            r.font.color.rgb = ACCENT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            # bottom border for scan-ability (still single-column, ATS-safe)
            _bottom_border(p)
        elif kind == "entry":
            p = doc.add_paragraph()
            r = p.add_run(_strip_md_inline(val))
            r.bold = True
            r.font.size = Pt(ENTRY_SIZE)
            r.font.color.rgb = DARK
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
        elif kind == "bullet":
            level, text = val
            style = "List Bullet" if level == 1 else "List Bullet 2"
            try:
                p = doc.add_paragraph(_strip_md_inline(text), style=style)
            except KeyError:
                p = doc.add_paragraph(_strip_md_inline(text), style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
        elif kind == "para":
            p = doc.add_paragraph()
            r = p.add_run(_strip_md_inline(val))
            r.font.size = Pt(BODY_SIZE)
            r.italic = _looks_like_dates(val)
            p.paragraph_format.space_after = Pt(2)
    doc.save(out_path)


def _looks_like_dates(text):
    return bool(re.search(r"\b(19|20)\d{2}\b", text) and
                re.search(r"present|–|-|\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)", text, re.I))


def _bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pbdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "B0B0B0"})
    pbdr.append(bottom)
    pPr.append(pbdr)


def build_txt(tokens, out_path):
    out = []
    for kind, val in tokens:
        if kind == "name":
            out.append(_strip_md_inline(val).upper())
        elif kind == "contact":
            out.append(_strip_md_inline(val))
            out.append("")
        elif kind == "section":
            out.append("")
            out.append(_strip_md_inline(val).upper())
            out.append("=" * len(val))
        elif kind == "entry":
            out.append(_strip_md_inline(val))
        elif kind == "bullet":
            level, text = val
            prefix = "  - " if level == 1 else "    * "
            out.append(prefix + _strip_md_inline(text))
        elif kind == "para":
            out.append(_strip_md_inline(val))
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(out).rstrip() + "\n")


def main():
    ap = argparse.ArgumentParser(description="Render ATS-safe .docx + .txt from a markdown CV.")
    ap.add_argument("--in", dest="infile", required=True)
    ap.add_argument("--out-docx")
    ap.add_argument("--out-txt")
    ap.add_argument("--outdir")
    ap.add_argument("--basename", default=None,
                    help="output stem inside --outdir. Defaults to the INPUT file's own stem, so "
                         "CV.md -> CV.docx and CoverLetter.md -> CoverLetter.docx. (It used to "
                         "default to 'CV', which silently overwrote the CV when rendering a letter "
                         "into the same folder.)")
    ap.add_argument("--page", choices=("a4", "letter"), default="a4",
                    help="page size: a4 (UK/world, default) or letter (Canada/US market=ca).")
    args = ap.parse_args()

    with open(args.infile, encoding="utf-8") as f:
        md = f.read()
    tokens = parse(md)

    if args.outdir:
        os.makedirs(args.outdir, exist_ok=True)
        # Derive the stem from the input unless the caller named one explicitly. Rendering
        # CoverLetter.md into a job folder must never land on CV.docx.
        stem = args.basename or os.path.splitext(os.path.basename(args.infile))[0]
        out_docx = args.out_docx or os.path.join(args.outdir, stem + ".docx")
        out_txt = args.out_txt or os.path.join(args.outdir, stem + ".txt")
    else:
        out_docx = args.out_docx or (os.path.splitext(args.infile)[0] + ".docx")
        out_txt = args.out_txt or (os.path.splitext(args.infile)[0] + ".txt")

    build_docx(tokens, out_docx, page=args.page)
    build_txt(tokens, out_txt)
    print(f"Wrote {out_docx} ({args.page.upper()})")
    print(f"Wrote {out_txt}")


if __name__ == "__main__":
    main()
